from openai import OpenAI
from docx import Document
client = OpenAI()
# make prompt and input again but this time check the terminal before running
file_name = input("What should the name of the file be?")
article_origin = input("Enter where the artical is from")
prompt = """Write a summary of the following article from the {ARTICLE}, then create a Conservative political take, a Moderate political take, and a Liberal political take:

GOP presidential candidate Vivek Ramaswamy says the time for Americans to address systemic racism is long gone, and that racism has dwindled to near irrelevance in the U.S.

Ramaswamy made the comments during a campaign stop in Iowa on Friday, telling his audience that contemporary efforts to combat systemic racism are only worsening the issue. Rev. Samuel Ansong asked how the U.S. should grapple with the issue if not in the form of affirmative action and other similar programs.

"Was there a time and place for correcting for those past injustices? Yeah, it was like in 1870," Ramaswamy responded, according to the Des Moines Register.

Ramaswamy reportedly added that affirmative action and DEI programs were "anti-American at their core" and "inherently divisive.

HOUSE DEMOCRAT FROM MAINE RIPS STATE'S DECISION TO TAKE TRUMP OFF BALLOT

Vivek Ramaswamy at GOP presidential debate
GOP presidential candidate Vivek Ramaswamy says the time for American to address systemic racism is long gone, and that racism has dwindled to near irrelevance in the U.S. (Micah Green/Bloomberg via Getty Images)

"But at a certain point in time, I think that what you would think of as racism in this country — or you can fill in the blank, sexism or any other -ism or form of discrimination — it gets to be small enough, not to be zero, but small enough that the best thing we can do is let it atrophy to irrelevance," Ramaswamy added.

He went on to argue that current efforts to combat so-called systemic racism do more harm than good by drawing attention to race and highlighting divisions.

Vivek Ramaswamy in New Hampshire
Ramaswamy argues that current efforts to combat so-called systemic racism do more harm than good by drawing attention to race and highlighting divisions. (REUTERS/Brian Snyder)

The exchange comes as Ramaswamy and the other GOP 2024 hopefuls are in the home stretch for the presidential primary in Iowa on January 15.

BLACK VOTERS SAY THEY'RE TURNING AWAY FROM ‘WEAK’ BIDEN IN 2024: ‘HE DIDN’T CHANGE ANYTHING'

Former President Donald Trump holds a massive lead over his competitors in national polls, though candidates like Florida Gov. Ron DeSantis, former South Carolina Gov. Nikki Haley and Ramaswamy have invested a huge effort into swaying Iowa voters specifically.

Vivek Ramaswamy, Trump, DeSantis Haley
Entrepreneur and author Vivek Ramaswamy is hoping to separate himself from Republicans who have or have yet to jump into the GOP race including former President Trump, Florida Gov. Ron DeSantis, and former U.N. Ambassador Nikki Haley. (Getty Images)


CLICK HERE TO GET THE FOX NEWS APP

Nevertheless, a Fox Business poll of Iowa Republicans found that Trump has a lead of 34 points as of mid-December.

""".format(ARTICLE = article_origin)


completion = client.chat.completions.create(
  model="gpt-3.5-turbo",
  messages=[
    {"role": "system", "content": "You are a helpful research assistant for extemporaneous speakers in the NSDA."},
    {"role": "user", "content": str(prompt)},
    {"role": "assistant", "content": str(article_origin)}
  ]
)

text = str(completion.choices[0].message)

document = Document()
document.add_heading("Article from: {ARTICLE}".format(ARTICLE=article_origin))
document.add_paragraph(text)
document.save("{NAME}.docx".format(NAME=file_name))